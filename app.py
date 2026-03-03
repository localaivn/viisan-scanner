import streamlit as st
import subprocess
import os
from datetime import datetime
import time
import socket
import base64
from openai import OpenAI
from docx import Document as DocxDocument
from docx.shared import Pt
import cv2
import numpy as np
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas as pdf_canvas

# ─── Config ───────────────────────────────────────────────
def get_local_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
    finally:
        s.close()
    return ip

LOCAL_IP       = get_local_ip()
MJPEG_URL      = f"http://{LOCAL_IP}:8080"
DEVICE         = "/dev/video0"
CAPTURE_WIDTH  = 3264
CAPTURE_HEIGHT = 2448
SAVE_DIR       = "captures"
OCR_MAX_WIDTH  = 1600
OCR_MAX_HEIGHT = 1200
CSS_FILE       = "static/style.css"
ENV_FILE       = ".env"

os.makedirs(SAVE_DIR, exist_ok=True)

# ─── Page config ──────────────────────────────────────────
st.set_page_config(page_title="Scanner Preview", page_icon="📷", layout="wide")

# ─── Load CSS ─────────────────────────────────────────────
def load_css(path: str):
    with open(path, "r") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

load_css(CSS_FILE)


def load_settings(path: str) -> tuple[str, str, list[str]]:
    theme = "light"
    api_key = ""
    env_lines: list[str] = []

    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            env_lines = f.readlines()

        for raw_line in env_lines:
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            clean_value = value.strip().strip('"').strip("'")
            if key.strip() == "DASHSCOPE_API_KEY":
                api_key = clean_value
            elif key.strip() == "UI_THEME" and clean_value in {"light", "dark"}:
                theme = clean_value

    return theme, api_key, env_lines


def save_settings(path: str, theme: str, api_key: str, env_lines: list[str]):
    updates = {
        "UI_THEME": theme,
        "DASHSCOPE_API_KEY": api_key,
    }
    output = []
    seen = set()

    for raw_line in env_lines:
        stripped = raw_line.strip()
        if stripped and not stripped.startswith("#") and "=" in stripped:
            key = stripped.split("=", 1)[0].strip()
            if key in updates:
                output.append(f"{key}={updates[key]}\n")
                seen.add(key)
                continue
        output.append(raw_line)

    for key, value in updates.items():
        if key not in seen:
            output.append(f"{key}={value}\n")

    with open(path, "w", encoding="utf-8") as f:
        f.writelines(output)


def apply_theme(theme: str):
    modal_css = """
    [data-testid="stDialog"] [role="dialog"] {
        background: #ffffff !important;
        color: #111111 !important;
    }
    [data-testid="stDialog"] [role="dialog"] * {
        color: #111111 !important;
    }
    [data-testid="stDialog"] input,
    [data-testid="stDialog"] textarea,
    [data-testid="stDialog"] [data-baseweb="input"] > div {
        background: #ffffff !important;
        color: #111111 !important;
    }
    """

    if theme == "dark":
        css = f"""
        <style>
        .stApp {{background-color: #1e232a !important; color: #e6edf3 !important;}}
        .app-header h1, .preview-info, .info-card, .footer {{color: #e6edf3 !important;}}
        .info-card {{background: #2b313a !important; border-left-color: #63a4ff !important;}}
        .preview-box {{background: #2b313a !important;}}
        .preview-info code {{background: #11161d !important; color: #d7e2ee !important;}}
        {modal_css}
        </style>
        """
    else:
        css = f"""
        <style>
        .stApp {{background-color: #b0b8c1 !important; color: #1a1a2e !important;}}
        section[data-testid="stSidebar"] {{background-color: #9aa3ad !important;}}
        {modal_css}
        </style>
        """
    st.markdown(css, unsafe_allow_html=True)

# ─── PDF conversion ───────────────────────────────────────
def crop_document_border(img_bgr: np.ndarray) -> np.ndarray:
    """
    Tự động phát hiện viền giấy và crop sát nội dung.
    Dùng Canny + findContours để tìm contour lớn nhất (tờ giấy).
    Fallback về ảnh gốc nếu không tìm thấy.
    """
    orig = img_bgr.copy()
    h, w = img_bgr.shape[:2]

    # Resize nhỏ để xử lý nhanh
    scale  = 800 / max(h, w)
    small  = cv2.resize(img_bgr, (int(w * scale), int(h * scale)))
    gray   = cv2.cvtColor(small, cv2.COLOR_BGR2GRAY)
    blur   = cv2.GaussianBlur(gray, (5, 5), 0)
    edges  = cv2.Canny(blur, 50, 150)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
    edges  = cv2.dilate(edges, kernel)

    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not contours:
        return orig

    # Lấy contour lớn nhất
    largest = max(contours, key=cv2.contourArea)

    # Cần contour đủ lớn (> 20% diện tích ảnh)
    if cv2.contourArea(largest) < 0.20 * small.shape[0] * small.shape[1]:
        return orig

    # Xấp xỉ thành tứ giác
    peri   = cv2.arcLength(largest, True)
    approx = cv2.approxPolyDP(largest, 0.02 * peri, True)

    if len(approx) == 4:
        # Perspective transform
        pts = approx.reshape(4, 2).astype(np.float32) / scale

        # Sắp xếp: top-left, top-right, bottom-right, bottom-left
        s    = pts.sum(axis=1)
        diff = np.diff(pts, axis=1)
        rect = np.array([
            pts[np.argmin(s)],
            pts[np.argmin(diff)],
            pts[np.argmax(s)],
            pts[np.argmax(diff)],
        ], dtype=np.float32)

        wA = np.linalg.norm(rect[2] - rect[3])
        wB = np.linalg.norm(rect[1] - rect[0])
        hA = np.linalg.norm(rect[1] - rect[2])
        hB = np.linalg.norm(rect[0] - rect[3])
        dw = int(max(wA, wB))
        dh = int(max(hA, hB))

        if dw < 100 or dh < 100:
            return orig

        dst = np.array([
            [0, 0], [dw - 1, 0],
            [dw - 1, dh - 1], [0, dh - 1]
        ], dtype=np.float32)

        M       = cv2.getPerspectiveTransform(rect, dst)
        warped  = cv2.warpPerspective(orig, M, (dw, dh))
        return warped
    else:
        # Không phải tứ giác → dùng bounding rect
        x, y, bw, bh = cv2.boundingRect(largest)
        # Scale về ảnh gốc
        x  = int(x  / scale); y  = int(y  / scale)
        bw = int(bw / scale); bh = int(bh / scale)
        # Thêm padding nhỏ
        pad = 10
        x1 = max(0, x - pad);  y1 = max(0, y - pad)
        x2 = min(w, x + bw + pad); y2 = min(h, y + bh + pad)
        return orig[y1:y2, x1:x2]


def image_to_pdf(src_path: str, crop_border: bool = True) -> str:
    """
    Chuyển ảnh sang PDF khổ A4, căn giữa, giữ tỉ lệ.
    Nếu crop_border=True thì tự động cắt viền giấy trước.
    Trả về đường dẫn file PDF.
    """
    img_bgr = cv2.imread(src_path)

    if crop_border:
        img_bgr = crop_document_border(img_bgr)

    # BGR → RGB rồi sang PIL
    img_rgb = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(img_rgb)

    # Kích thước A4 theo point (1 inch = 72pt, A4 = 210×297mm)
    a4_w, a4_h = A4          # (595.27, 841.89) pt
    margin     = 28.35        # 10mm margin

    avail_w = a4_w - 2 * margin
    avail_h = a4_h - 2 * margin

    img_w, img_h = pil_img.size
    ratio  = min(avail_w / img_w, avail_h / img_h)
    draw_w = img_w * ratio
    draw_h = img_h * ratio

    # Căn giữa
    x_off = margin + (avail_w - draw_w) / 2
    y_off = margin + (avail_h - draw_h) / 2

    # Lưu ảnh tạm
    base_name = os.path.splitext(os.path.basename(src_path))[0]
    pdf_path  = f"{SAVE_DIR}/{base_name}.pdf"
    tmp_img   = f"{SAVE_DIR}/_tmp_pdf.jpg"
    pil_img.save(tmp_img, "JPEG", quality=92)

    # Tạo PDF
    c = pdf_canvas.Canvas(pdf_path, pagesize=A4)
    c.drawImage(tmp_img, x_off, y_off, width=draw_w, height=draw_h)
    c.save()

    os.remove(tmp_img)
    return pdf_path


# ─── Session state ────────────────────────────────────────
if "captured_file" not in st.session_state:
    st.session_state.captured_file = None
if "ocr_text" not in st.session_state:
    st.session_state.ocr_text = None
if "pdf_file" not in st.session_state:
    st.session_state.pdf_file = None
if "theme" not in st.session_state or "dashscope_api_key" not in st.session_state:
    loaded_theme, loaded_api_key, loaded_env_lines = load_settings(ENV_FILE)
    st.session_state.theme = loaded_theme
    st.session_state.dashscope_api_key = loaded_api_key
    st.session_state.env_lines = loaded_env_lines

apply_theme(st.session_state.theme)


@st.dialog("Settings")
def open_settings_dialog():
    theme_value = st.radio(
        "Giao diện",
        options=["light", "dark"],
        horizontal=True,
        format_func=lambda x: "Sáng" if x == "light" else "Tối",
        index=0 if st.session_state.theme == "light" else 1,
    )
    api_key_value = st.text_input(
        "🔑 DashScope API Key",
        type="password",
        value=st.session_state.dashscope_api_key,
        placeholder="sk-...",
    )

    if st.button("💾 Lưu settings"):
        save_settings(
            ENV_FILE,
            theme=theme_value,
            api_key=api_key_value.strip(),
            env_lines=st.session_state.env_lines,
        )
        updated_theme, updated_api_key, updated_env_lines = load_settings(ENV_FILE)
        st.session_state.theme = updated_theme
        st.session_state.dashscope_api_key = updated_api_key
        st.session_state.env_lines = updated_env_lines
        st.success("Đã lưu settings vào file .env")
        st.rerun()

# ─── Header ───────────────────────────────────────────────
header_left, header_right = st.columns([12, 2], vertical_alignment="center")
with header_left:
    st.markdown("""
    <div class="app-header">
        <span class="app-header-icon">📷</span>
        <h1>Scanner Preview &amp; Capture</h1>
        <span class="live-badge"><span class="live-dot"></span>LIVE</span>
    </div>
    """, unsafe_allow_html=True)
with header_right:
    if st.button("⚙️ Setting"):
        open_settings_dialog()

# ─── Layout chính ─────────────────────────────────────────
col_preview, col_ctrl = st.columns([3, 1], gap="medium")

with col_preview:
    st.markdown(f"""
    <div class="preview-box">
        <img src="{MJPEG_URL}" alt="Live Preview">
    </div>
    <div class="preview-info">
        🌐 <code>{MJPEG_URL}</code> &nbsp;|&nbsp; 📹 <code>{DEVICE}</code>
    </div>
    """, unsafe_allow_html=True)

with col_ctrl:
    # ── Thông số ──
    st.markdown(f"""
    <div class="info-card">
        📐 Capture: <b>{CAPTURE_WIDTH} × {CAPTURE_HEIGHT}</b> px<br>
        🔍 OCR input: <b>{OCR_MAX_WIDTH} × {OCR_MAX_HEIGHT}</b> px<br>
        💾 Save: <b>{SAVE_DIR}/</b>
    </div>
    """, unsafe_allow_html=True)

    # ── Exposure ──
    exposure = st.slider(
        "☀️ Exposure", min_value=100, max_value=20000,
        value=12000, step=100, help="Cao hơn = sáng hơn"
    )

    st.markdown("<div style='margin-top:6px'></div>", unsafe_allow_html=True)

    # ── Nút Capture + Convert PDF cùng hàng ──
    btn_col1, btn_col2 = st.columns(2, gap="small")
    with btn_col1:
        capture_btn = st.button("📸 Capture")
    with btn_col2:
        crop_border = st.checkbox("✂️ Crop border", value=True,
                                  help="Tự động cắt viền & căn phối cảnh tờ giấy")
        pdf_disabled = (
            st.session_state.captured_file is None or
            not os.path.exists(st.session_state.captured_file or "")
        )
        pdf_btn = st.button("📄 Convert to PDF", disabled=pdf_disabled)

    st.markdown("<div style='margin-top:2px'></div>", unsafe_allow_html=True)

    # ── Download ảnh ──
    if st.session_state.captured_file and os.path.exists(st.session_state.captured_file):
        fname = os.path.basename(st.session_state.captured_file)
        with open(st.session_state.captured_file, "rb") as f:
            st.download_button("⬇️ Download Image", f,
                               file_name=fname, mime="image/jpeg")

    # ── Download PDF ──
    if st.session_state.pdf_file and os.path.exists(st.session_state.pdf_file):
        with open(st.session_state.pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF", f,
                               file_name=os.path.basename(st.session_state.pdf_file),
                               mime="application/pdf")

    # ── Nút OCR ──
    ocr_disabled = (
        st.session_state.captured_file is None or
        not os.path.exists(st.session_state.captured_file or "")
    )
    ocr_btn = st.button("🔍 Run OCR → .docx", disabled=ocr_disabled)

    # ── OCR result + Download docx ──
    if st.session_state.ocr_text:
        st.markdown("<div style='margin-top:4px'></div>", unsafe_allow_html=True)
        edited_text = st.text_area(
            "📄 OCR Result",
            value=st.session_state.ocr_text,
            height=280,
        )

        base_name = os.path.splitext(
            os.path.basename(st.session_state.captured_file)
        )[0]
        docx_path = f"{SAVE_DIR}/{base_name}_ocr.docx"

        doc = DocxDocument()
        doc.add_heading("OCR Result", level=1)
        doc.add_paragraph(f"Source: {os.path.basename(st.session_state.captured_file)}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")
        for line in edited_text.split("\n"):
            para = doc.add_paragraph(line)
            para.style.font.size = Pt(12)
        doc.save(docx_path)

        with open(docx_path, "rb") as f:
            st.download_button(
                "⬇️ Download .docx", f,
                file_name=os.path.basename(docx_path),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # ── Footer ──
    st.markdown("""
    <div class="footer">
        Scanner Preview v1.3 &nbsp;|&nbsp;
        <a href="https://ttaisolutions.com" target="_blank">TTAI Solutions Software</a>
    </div>
    """, unsafe_allow_html=True)

# ─── Capture logic ────────────────────────────────────────
if capture_btn:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"{SAVE_DIR}/capture_{timestamp}.jpg"

    prog = st.progress(0, text="Stopping MJPEG service...")
    subprocess.run(["sudo", "systemctl", "stop", "mjpeg-server.service"],
                   capture_output=True)
    time.sleep(0.3)

    prog.progress(20, text="Setting exposure...")
    subprocess.run(["v4l2-ctl", "-d", DEVICE, "--set-ctrl=exposure_auto=1"],
                   capture_output=True)
    subprocess.run(["v4l2-ctl", "-d", DEVICE,
                    f"--set-ctrl=exposure_absolute={exposure}"],
                   capture_output=True)

    prog.progress(40, text="Waiting sensor stabilize...")
    time.sleep(1.0)

    prog.progress(60, text="Capturing...")
    result = subprocess.run([
        "ffmpeg", "-y",
        "-f", "v4l2", "-framerate", "3",
        "-input_format", "mjpeg",
        "-video_size", f"{CAPTURE_WIDTH}x{CAPTURE_HEIGHT}",
        "-i", DEVICE,
        "-vf", "transpose=1",
        "-frames:v", "3", "-update", "1",
        "-q:v", "2", filename
    ], stdout=subprocess.PIPE, stderr=subprocess.PIPE)

    prog.progress(85, text="Restarting MJPEG service...")
    subprocess.run(["sudo", "systemctl", "start", "mjpeg-server.service"],
                   capture_output=True)

    prog.progress(100, text="Done!")
    time.sleep(0.3)
    prog.empty()

    if result.returncode == 0 and os.path.exists(filename):
        st.session_state.captured_file = filename
        st.session_state.ocr_text = None
        st.session_state.pdf_file = None
        st.success(f"✅ Saved: `{os.path.basename(filename)}`")
        st.rerun()
    else:
        st.error("❌ Capture failed!")
        with st.expander("FFmpeg log"):
            st.code(result.stderr.decode(), language="bash")

# ─── Convert to PDF logic ─────────────────────────────────
if pdf_btn:
    with st.spinner("📄 Converting to PDF..."):
        try:
            pdf_path = image_to_pdf(
                st.session_state.captured_file,
                crop_border=crop_border
            )
            st.session_state.pdf_file = pdf_path
            st.success(f"✅ PDF: `{os.path.basename(pdf_path)}`")
            st.rerun()
        except Exception as e:
            st.error(f"❌ Convert failed: {e}")

# ─── OCR logic ────────────────────────────────────────────
if ocr_btn:
    api_key = st.session_state.dashscope_api_key
    if not api_key:
        st.warning("⚠️ Nhập API Key trước khi chạy OCR.")
    else:
        with st.spinner("🔄 Resizing image..."):
            img = cv2.imread(st.session_state.captured_file)
            h, w = img.shape[:2]
            scale = min(OCR_MAX_WIDTH / w, OCR_MAX_HEIGHT / h, 1.0)
            if scale < 1.0:
                img = cv2.resize(img, (int(w * scale), int(h * scale)),
                                 interpolation=cv2.INTER_AREA)
            _, buf = cv2.imencode('.jpg', img, [cv2.IMWRITE_JPEG_QUALITY, 90])
            img_b64 = base64.b64encode(buf.tobytes()).decode()

        with st.spinner("🤖 Running OCR..."):
            try:
                client = OpenAI(
                    api_key=api_key,
                    base_url="https://dashscope-intl.aliyuncs.com/compatible-mode/v1"
                )
                resp = client.chat.completions.create(
                    model="qwen-vl-ocr-2025-11-20",
                    messages=[{
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
                                "min_pixels": 100 * 100,
                                "max_pixels": OCR_MAX_WIDTH * OCR_MAX_HEIGHT
                            },
                            {
                                "type": "text",
                                "text": "Read all the text in this image. Output the text exactly as it appears, preserving layout and line breaks."
                            }
                        ]
                    }],
                    max_tokens=4096
                )
                st.session_state.ocr_text = resp.choices[0].message.content
                st.rerun()
            except Exception as e:
                st.error(f"❌ OCR failed: {e}")
